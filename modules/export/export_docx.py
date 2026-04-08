from docx import Document
from docx.shared import Pt

def exportar_informe_docx(informe, nombre_archivo):

    doc = Document()

    # Membrete
    doc.add_heading("MUNICIPALIDAD PROVINCIAL DE CALCA", level=1)
    doc.add_paragraph("OFICINA DE ASESORÍA JURÍDICA")
    doc.add_paragraph("")

    doc.add_heading("INFORME LEGAL", level=2)

    doc.add_paragraph(informe)

    ruta = f"data/{nombre_archivo}.docx"
    doc.save(ruta)

    return ruta
